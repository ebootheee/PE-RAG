#!/usr/bin/env python3
"""
Interactive Legal Document Vector Loader
User-friendly version with interactive setup and guided configuration.

This script helps you process legal documents and upload them to a Pinecone vector database
for semantic search and retrieval. It will guide you through the setup process step by step.

Features:
- Interactive API key setup and validation
- Connection testing
- Customizable processing parameters with recommendations
- Real-time progress tracking
- Comprehensive error handling

Usage:
    python vector_loader.py
"""

import os
import re
import json
import logging
import hashlib
import asyncio
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Optional, Generator
import argparse
from dataclasses import dataclass
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor, as_completed
import threading
from queue import Queue
import time
import getpass
import sys

# Core libraries
import fitz  # PyMuPDF
from docx import Document as DocxDocument
import mammoth  # For .doc files
from tqdm import tqdm

# LangChain and vector processing
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

# Vector and embedding services
from pinecone import Pinecone, ServerlessSpec
from openai import OpenAI

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('vector_loader.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Configuration
DEFAULT_DOCS_PATH = r"C:\Documents\Legal_Documents"
PINECONE_INDEX_NAME = "legal-documents"
PINECONE_NAMESPACE = "legal-docs"
OPENAI_EMBEDDING_MODEL_LARGE = "text-embedding-3-large"
OPENAI_EMBEDDING_MODEL_SMALL = "text-embedding-3-small"
EMBEDDING_DIMENSIONS_LARGE = 3072
EMBEDDING_DIMENSIONS_SMALL = 1536
CHUNK_SIZE = 2000
CHUNK_OVERLAP = 400
BATCH_SIZE = 50
MAX_WORKERS = 4
EMBEDDING_BATCH_SIZE = 25
MAX_RETRIES = 5
RATE_LIMIT_DELAY = 1.0
MAX_TOKENS_PER_MIN = 950000

def print_banner():
    """Print welcome banner"""
    print("\n" + "="*80)
    print("🚀 INTERACTIVE LEGAL DOCUMENT VECTOR LOADER")
    print("="*80)
    print("This tool will help you process legal documents and upload them to Pinecone")
    print("for semantic search and AI-powered document retrieval.")
    print("="*80 + "\n")

def get_api_keys():
    """Interactively collect API keys"""
    print("📋 API KEY SETUP")
    print("-" * 20)
    print("You'll need API keys for both OpenAI and Pinecone services.\n")
    
    # Get OpenAI API key
    openai_key = os.getenv('OPENAI_API_KEY')
    if openai_key:
        use_existing = input(f"🔑 Found existing OpenAI API key (ends with ...{openai_key[-8:]}). Use it? (y/n): ").lower().strip()
        if use_existing != 'y':
            openai_key = None
    
    if not openai_key:
        print("\n🔑 Enter your OpenAI API key:")
        print("   (Get one at: https://platform.openai.com/api-keys)")
        openai_key = getpass.getpass("   OpenAI API Key: ").strip()
    
    # Get Pinecone API key
    pinecone_key = os.getenv('PINECONE_API_KEY')
    if pinecone_key:
        use_existing = input(f"\n🔑 Found existing Pinecone API key (ends with ...{pinecone_key[-8:]}). Use it? (y/n): ").lower().strip()
        if use_existing != 'y':
            pinecone_key = None
    
    if not pinecone_key:
        print("\n🔑 Enter your Pinecone API key:")
        print("   (Get one at: https://www.pinecone.io/)")
        pinecone_key = getpass.getpass("   Pinecone API Key: ").strip()
    
    return openai_key, pinecone_key

def test_connections(openai_key: str, pinecone_key: str):
    """Test API connections"""
    print("\n🔄 TESTING CONNECTIONS")
    print("-" * 25)
    
    # Test OpenAI
    try:
        print("   Testing OpenAI connection...", end=" ")
        openai_client = OpenAI(api_key=openai_key)
        # Test with a small request
        response = openai_client.embeddings.create(
            model="text-embedding-3-large",
            input=["test"],
            dimensions=3072
        )
        print("✅ Success!")
    except Exception as e:
        print(f"❌ Failed: {str(e)}")
        return False
    
    # Test Pinecone
    try:
        print("   Testing Pinecone connection...", end=" ")
        pc = Pinecone(api_key=pinecone_key)
        # List indexes to test connection
        indexes = pc.list_indexes()
        print("✅ Success!")
        
        # Show available indexes
        if hasattr(indexes, 'indexes') and indexes.indexes:
            print(f"   Found {len(indexes.indexes)} index(es):")
            for idx in indexes.indexes:
                print(f"     - {idx.name}")
        
    except Exception as e:
        print(f"❌ Failed: {str(e)}")
        return False
    
    print("\n✅ All connections successful!")
    return True

def get_processing_parameters():
    """Interactively collect processing parameters"""
    print("\n⚙️  PROCESSING PARAMETERS")
    print("-" * 30)
    print("Configure how documents will be processed and uploaded.\n")
    
    # Batch size
    print("📦 Batch Size (number of vectors uploaded at once)")
    print("   Recommendation: 50 (good balance of speed and reliability)")
    batch_size = input("   Enter batch size [50]: ").strip()
    batch_size = int(batch_size) if batch_size.isdigit() else 50
    
    # Chunk size
    print("\n📄 Chunk Size (characters per document chunk)")
    print("   Recommendation: 2000 (optimal for legal documents)")
    print("   Smaller = more precise, Larger = more context")
    chunk_size = input("   Enter chunk size [2000]: ").strip()
    chunk_size = int(chunk_size) if chunk_size.isdigit() else 2000
    
    # Overlap
    print("\n🔗 Chunk Overlap (character overlap between chunks)")
    print("   Recommendation: 400 (maintains context continuity)")
    overlap = input("   Enter overlap [400]: ").strip()
    overlap = int(overlap) if overlap.isdigit() else 400
    
    # Workers
    print("\n👥 Worker Count (parallel processing threads)")
    print("   Recommendation: 4 (good balance, avoids rate limits)")
    print("   More workers = faster but may hit API limits")
    workers = input("   Enter worker count [4]: ").strip()
    workers = int(workers) if workers.isdigit() else 4
    
    return batch_size, chunk_size, overlap, workers

def get_storage_parameters():
    """Interactively collect storage parameters"""
    print("\n💾 STORAGE PARAMETERS")
    print("-" * 25)
    
    # Documents folder
    print(f"📁 Documents Folder")
    print(f"   Default: {DEFAULT_DOCS_PATH}")
    docs_path = input(f"   Enter documents folder path [{DEFAULT_DOCS_PATH}]: ").strip()
    docs_path = docs_path if docs_path else DEFAULT_DOCS_PATH
    
    # Validate path
    if not os.path.exists(docs_path):
        print(f"   ⚠️  Warning: Path does not exist: {docs_path}")
        create_path = input("   Continue anyway? (y/n): ").lower().strip()
        if create_path != 'y':
            return None, None, None
    
    # Namespace
    print(f"\n📂 Pinecone Namespace (logical separation of vectors)")
    print(f"   Default: {PINECONE_NAMESPACE}")
    namespace = input(f"   Enter namespace [{PINECONE_NAMESPACE}]: ").strip()
    namespace = namespace if namespace else PINECONE_NAMESPACE
    
    # Index name
    print(f"\n🗃️  Pinecone Index Name")
    print(f"   Default: {PINECONE_INDEX_NAME}")
    index_name = input(f"   Enter index name [{PINECONE_INDEX_NAME}]: ").strip()
    index_name = index_name if index_name else PINECONE_INDEX_NAME
    
    return docs_path, namespace, index_name

def get_embedding_parameters():
    """Interactively collect embedding parameters"""
    print("\n🧠 EMBEDDING PARAMETERS")
    print("-" * 28)
    
    # Model selection
    print("🤖 Embedding Model")
    print("   1. Large (text-embedding-3-large) - Higher quality, 3072 dimensions")
    print("   2. Small (text-embedding-3-small) - Faster, cheaper, 1536 dimensions")
    print("   Recommendation: Large (better for legal documents)")
    
    while True:
        choice = input("   Enter choice (1 for Large, 2 for Small) [1]: ").strip()
        if choice == '2':
            model = OPENAI_EMBEDDING_MODEL_SMALL
            dimensions = EMBEDDING_DIMENSIONS_SMALL
            break
        elif choice == '1' or choice == '':
            model = OPENAI_EMBEDDING_MODEL_LARGE
            dimensions = EMBEDDING_DIMENSIONS_LARGE
            break
        else:
            print("   Please enter 1 or 2")
    
    # Dimensions (allow override)
    print(f"\n📏 Vector Dimensions")
    print(f"   Default for {model}: {dimensions}")
    custom_dims = input(f"   Enter custom dimensions [{dimensions}]: ").strip()
    dimensions = int(custom_dims) if custom_dims.isdigit() else dimensions
    
    return model, dimensions

def show_configuration_summary(config):
    """Display final configuration summary"""
    print("\n📋 CONFIGURATION SUMMARY")
    print("=" * 40)
    print(f"📁 Documents Folder:    {config['docs_path']}")
    print(f"🗃️  Pinecone Index:      {config['index_name']}")
    print(f"📂 Namespace:           {config['namespace']}")
    print(f"🤖 Embedding Model:     {config['model']}")
    print(f"📏 Dimensions:          {config['dimensions']}")
    print(f"📦 Batch Size:          {config['batch_size']}")
    print(f"📄 Chunk Size:          {config['chunk_size']}")
    print(f"🔗 Overlap:             {config['overlap']}")
    print(f"👥 Workers:             {config['workers']}")
    print("=" * 40)
    
    # Count documents
    if os.path.exists(config['docs_path']):
        doc_count = len([f for f in Path(config['docs_path']).rglob('*') 
                        if f.suffix.lower() in ['.pdf', '.docx', '.doc', '.txt']])
        print(f"📊 Documents found:     {doc_count}")
    
    print("\nReady to process documents!")
    confirm = input("\n🚀 Start processing? (y/n): ").lower().strip()
    return confirm == 'y'

@dataclass
class ProcessingStats:
    """Thread-safe processing statistics"""
    def __init__(self):
        self.lock = threading.Lock()
        self.total_files: int = 0
        self.processed_files: int = 0
        self.failed_files: int = 0
        self.total_chunks: int = 0
        self.uploaded_vectors: int = 0
        self.skipped_files: int = 0
        self.processing_time: float = 0.0
    
    def increment_processed(self, chunks_count: int = 0):
        with self.lock:
            self.processed_files += 1
            self.total_chunks += chunks_count
    
    def increment_failed(self):
        with self.lock:
            self.failed_files += 1
    
    def add_uploaded(self, count: int):
        with self.lock:
            self.uploaded_vectors += count

class ParallelLegalDocumentVectorLoader:
    def __init__(self, 
                 docs_directory: str,
                 pinecone_api_key: str,
                 openai_api_key: str,
                 index_name: str = PINECONE_INDEX_NAME,
                 namespace: str = PINECONE_NAMESPACE,
                 max_workers: int = MAX_WORKERS,
                 embedding_model: str = None,
                 embedding_dimensions: int = None):
        
        self.docs_directory = Path(docs_directory)
        self.index_name = index_name
        self.namespace = namespace
        self.max_workers = max_workers
        self.stats = ProcessingStats()
        
        # Use provided embedding parameters or defaults
        self.embedding_model = embedding_model or OPENAI_EMBEDDING_MODEL_LARGE
        self.embedding_dimensions = embedding_dimensions or EMBEDDING_DIMENSIONS_LARGE
        
        # Initialize services
        self.pinecone_client = Pinecone(api_key=pinecone_api_key)
        self.openai_client = OpenAI(api_key=openai_api_key)
        
        # Text splitter optimized for legal documents
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP,
            separators=[
                "ARTICLE ",
                "SECTION ", 
                "§ ",
                "(a) ", "(b) ", "(c) ", "(d) ",
                "WHEREAS",
                "NOW, THEREFORE",
                "DEFINITIONS",
                "EXHIBIT ",
                "SCHEDULE ",
                "\n\n",
                "\n",
                " "
            ],
            length_function=len,
            keep_separator=True
        )
        
        # Supported file extensions
        self.supported_extensions = {'.pdf', '.docx', '.doc', '.txt'}
        
        # Vector upload queue for background processing
        self.upload_queue = Queue()
        self.upload_thread = None
        self.stop_upload_thread = False

    def setup_pinecone_index(self) -> bool:
        """Create Pinecone index if it doesn't exist"""
        try:
            # Check if index exists
            existing_indexes = [index.name for index in self.pinecone_client.list_indexes()]
            
            if self.index_name not in existing_indexes:
                logger.info(f"Creating Pinecone index: {self.index_name}")
                self.pinecone_client.create_index(
                    name=self.index_name,
                    dimension=EMBEDDING_DIMENSIONS,
                    metric="cosine",
                    spec=ServerlessSpec(
                        cloud="aws",
                        region="us-east-1"
                    )
                )
                logger.info("Index created successfully")
            else:
                logger.info(f"Index {self.index_name} already exists")
            
            # Get index reference
            self.index = self.pinecone_client.Index(self.index_name)
            return True
            
        except Exception as e:
            logger.error(f"Failed to setup Pinecone index: {str(e)}")
            return False

    def _start_background_uploader(self):
        """Start background thread for uploading vectors"""
        def upload_worker():
            batch_vectors = []
            
            while not self.stop_upload_thread or not self.upload_queue.empty():
                try:
                    # Get vectors from queue with timeout
                    try:
                        vectors = self.upload_queue.get(timeout=1.0)
                        if vectors is None:  # Poison pill
                            break
                        batch_vectors.extend(vectors)
                    except:
                        continue
                    
                    # Upload when batch is full or queue is empty
                    if len(batch_vectors) >= BATCH_SIZE or (self.upload_queue.empty() and batch_vectors):
                        upload_batch = batch_vectors[:BATCH_SIZE]
                        batch_vectors = batch_vectors[BATCH_SIZE:]
                        
                        self.index.upsert(
                            vectors=upload_batch,
                            namespace=self.namespace
                        )
                        
                        self.stats.add_uploaded(len(upload_batch))
                        logger.debug(f"Uploaded batch of {len(upload_batch)} vectors")
                
                except Exception as e:
                    logger.error(f"Error in upload worker: {str(e)}")
            
            # Upload remaining vectors
            if batch_vectors:
                self.index.upsert(
                    vectors=batch_vectors,
                    namespace=self.namespace
                )
                self.stats.add_uploaded(len(batch_vectors))
        
        self.upload_thread = threading.Thread(target=upload_worker, daemon=True)
        self.upload_thread.start()

    def _stop_background_uploader(self):
        """Stop background uploader thread"""
        self.stop_upload_thread = True
        self.upload_queue.put(None)  # Poison pill
        if self.upload_thread:
            self.upload_thread.join()

    def scan_documents(self) -> List[Path]:
        """Scan directory for supported document files"""
        if not self.docs_directory.exists():
            raise FileNotFoundError(f"Documents directory not found: {self.docs_directory}")
        
        logger.info(f"Scanning documents in: {self.docs_directory}")
        
        files = []
        for file_path in self.docs_directory.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                files.append(file_path)
        
        self.stats.total_files = len(files)
        logger.info(f"Found {len(files)} supported documents")
        
        return sorted(files)  # Sort for consistent processing order

    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF using PyMuPDF"""
        try:
            doc = fitz.open(file_path)
            text = ""
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                text += page.get_text()
                text += "\n\n"  # Add page breaks
            
            doc.close()
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting PDF text from {file_path}: {str(e)}")
            raise

    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX"""
        try:
            doc = DocxDocument(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text from {file_path}: {str(e)}")
            raise

    def _extract_doc_text(self, file_path: Path) -> str:
        """Extract text from DOC using mammoth"""
        try:
            with open(file_path, "rb") as docx_file:
                result = mammoth.extract_raw_text(docx_file)
                return result.value.strip()
                
        except Exception as e:
            logger.error(f"Error extracting DOC text from {file_path}: {str(e)}")
            raise

    def _extract_txt_text(self, file_path: Path) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read().strip()
                
        except Exception as e:
            logger.error(f"Error extracting TXT text from {file_path}: {str(e)}")
            raise

    def _extract_metadata_from_filename(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from organized filename"""
        filename = file_path.name
        
        metadata = {
            'source_file': filename,
            'file_path': str(file_path),
            'file_size': file_path.stat().st_size,
            'file_extension': file_path.suffix.lower(),
            'processing_date': datetime.now().isoformat()
        }
        
        # Parse priority prefix (1_HIGH_, 2_MED_, 3_LOW_)
        priority_match = re.match(r'^(\d+)_([^_]+)_', filename)
        if priority_match:
            priority_num, priority_name = priority_match.groups()
            metadata['priority_number'] = int(priority_num)
            metadata['priority'] = priority_name.lower()
        
        # Parse document type prefix (PA_, SUB_, SL_, etc.)
        doc_type_match = re.search(r'_([A-Z]+)_', filename)
        if doc_type_match:
            doc_prefix = doc_type_match.group(1)
            metadata['document_type_prefix'] = doc_prefix
            metadata['document_type'] = self._map_doc_prefix_to_type(doc_prefix)
        
        # Extract organization information (generic placeholder)
        # Add custom logic here to identify document organization/source if needed
        
        # Check execution status
        metadata['is_executed'] = 'executed' in filename.lower()
        metadata['is_draft'] = 'draft' in filename.lower()
        metadata['is_redacted'] = 'redacted' in filename.lower()
        
        return metadata

    def _map_doc_prefix_to_type(self, prefix: str) -> str:
        """Map document prefix to full type name"""
        mapping = {
            'PA': 'partnership_agreement',
            'SUB': 'subscription_agreement', 
            'SL': 'side_letter',
            'MGMT': 'management_agreement',
            'AMD': 'amendment',
            'CONS': 'consent',
            'OPN': 'opinion',
            'PPM': 'private_placement_memorandum',
            'SCH': 'schedule',
            'CERT': 'certificate',
            'FORMD': 'form_d',
            'ASSGN': 'assignment',
            'GUAR': 'guarantee',
            'OA': 'operating_agreement',
            'BYL': 'bylaws'
        }
        return mapping.get(prefix, 'legal_document')

    def process_single_document(self, file_path: Path) -> List[Document]:
        """Process a single document into chunks"""
        try:
            logger.debug(f"Processing: {file_path.name}")
            
            # Extract text using appropriate handler
            file_handlers = {
                '.pdf': self._extract_pdf_text,
                '.docx': self._extract_docx_text,
                '.doc': self._extract_doc_text,
                '.txt': self._extract_txt_text
            }
            
            handler = file_handlers.get(file_path.suffix.lower())
            if not handler:
                raise ValueError(f"No handler for file type: {file_path.suffix}")
            
            text_content = handler(file_path)
            
            if not text_content or len(text_content.strip()) < 50:
                raise ValueError("Document appears to be empty or too short")
            
            # Extract metadata
            base_metadata = self._extract_metadata_from_filename(file_path)
            
            # Split text into chunks
            chunks = self.text_splitter.split_text(text_content)
            
            # Create Document objects with metadata
            documents = []
            for i, chunk in enumerate(chunks):
                chunk_metadata = base_metadata.copy()
                chunk_metadata.update({
                    'chunk_index': i,
                    'total_chunks': len(chunks),
                    'chunk_id': self._generate_chunk_id(file_path, i),
                    'chunk_size': len(chunk)
                })
                
                documents.append(Document(
                    page_content=chunk,
                    metadata=chunk_metadata
                ))
            
            return documents
            
        except Exception as e:
            logger.error(f"Failed to process {file_path}: {str(e)}")
            raise

    def _generate_chunk_id(self, file_path: Path, chunk_index: int) -> str:
        """Generate unique chunk ID"""
        file_hash = hashlib.md5(str(file_path).encode()).hexdigest()[:8]
        return f"{file_hash}_{chunk_index:04d}"

    def create_embeddings_batch(self, documents: List[Document]) -> List[Dict[str, Any]]:
        """Create embeddings for a batch of document chunks"""
        if not documents:
            return []
        
        try:
            # Extract text content
            texts = [doc.page_content for doc in documents]
            
            # Create embeddings in smaller batches for API efficiency
            all_vectors = []
            
            for i in range(0, len(texts), EMBEDDING_BATCH_SIZE):
                batch_texts = texts[i:i + EMBEDDING_BATCH_SIZE]
                batch_docs = documents[i:i + EMBEDDING_BATCH_SIZE]
                
                response = self.openai_client.embeddings.create(
                    model=self.embedding_model,
                    input=batch_texts,
                    dimensions=self.embedding_dimensions
                )
                
                # Prepare vectors for Pinecone
                for doc, embedding_data in zip(batch_docs, response.data):
                    vector = {
                        'id': doc.metadata['chunk_id'],
                        'values': embedding_data.embedding,
                        'metadata': {
                            # Store text content in the standard field that LangChain integrations expect
                            'text': doc.page_content,  # Standard LangChain field for content
                            # Additional metadata (keep string/number types only)
                            'source_file': doc.metadata['source_file'],
                            'document_type': doc.metadata.get('document_type', 'legal_document'),
                            'priority': doc.metadata.get('priority', 'medium'),
                            'chunk_index': doc.metadata['chunk_index'],
                            'total_chunks': doc.metadata['total_chunks'],
                            'is_executed': doc.metadata.get('is_executed', False),
                            'is_draft': doc.metadata.get('is_draft', False),
                            'file_size': doc.metadata['file_size'],
                            'processing_date': doc.metadata['processing_date']
                        }
                    }
                    all_vectors.append(vector)
            
            return all_vectors
            
        except Exception as e:
            logger.error(f"Failed to create embeddings: {str(e)}")
            raise

    def process_document_worker(self, file_path: Path) -> Optional[List[Dict[str, Any]]]:
        """Worker function for processing a single document"""
        try:
            # Process document into chunks
            documents = self.process_single_document(file_path)
            
            if not documents:
                return None
            
            # Create embeddings
            vectors = self.create_embeddings_batch(documents)
            
            # Update stats
            self.stats.increment_processed(len(documents))
            
            return vectors
            
        except Exception as e:
            logger.error(f"Error processing {file_path}: {str(e)}")
            self.stats.increment_failed()
            return None

    def process_all_documents_parallel(self, batch_size: int = BATCH_SIZE) -> ProcessingStats:
        """Process all documents in parallel and load them into Pinecone"""
        start_time = datetime.now()
        
        try:
            # Setup Pinecone
            if not self.setup_pinecone_index():
                raise RuntimeError("Failed to setup Pinecone index")
            
            # Start background uploader
            self._start_background_uploader()
            
            # Scan for documents
            files = self.scan_documents()
            
            if not files:
                logger.warning("No documents found to process")
                return self.stats
            
            logger.info(f"Starting parallel processing of {len(files)} documents with {self.max_workers} workers")
            
            # Process documents in parallel
            with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
                # Submit all tasks
                future_to_file = {
                    executor.submit(self.process_document_worker, file_path): file_path 
                    for file_path in files
                }
                
                # Process completed tasks with progress bar
                with tqdm(total=len(files), desc="Processing documents") as pbar:
                    for future in as_completed(future_to_file):
                        file_path = future_to_file[future]
                        
                        try:
                            vectors = future.result()
                            
                            if vectors:
                                # Queue vectors for background upload
                                self.upload_queue.put(vectors)
                            
                        except Exception as e:
                            logger.error(f"Error processing {file_path}: {str(e)}")
                            self.stats.increment_failed()
                        
                        # Update progress bar
                        pbar.set_postfix({
                            'Processed': self.stats.processed_files,
                            'Failed': self.stats.failed_files,
                            'Chunks': self.stats.total_chunks,
                            'Uploaded': self.stats.uploaded_vectors
                        })
                        pbar.update(1)
            
            # Wait for all uploads to complete
            logger.info("Waiting for all uploads to complete...")
            while not self.upload_queue.empty():
                time.sleep(0.5)
            
            # Stop background uploader
            self._stop_background_uploader()
            
            # Calculate final stats
            end_time = datetime.now()
            self.stats.processing_time = (end_time - start_time).total_seconds()
            
            logger.info("Document processing completed successfully")
            self.print_final_stats()
            
            return self.stats
            
        except Exception as e:
            self._stop_background_uploader()
            logger.error(f"Document processing failed: {str(e)}")
            raise

    def test_connections(self) -> bool:
        """Test connections to Pinecone and OpenAI"""
        try:
            # Test Pinecone
            logger.info("Testing Pinecone connection...")
            indexes = self.pinecone_client.list_indexes()
            logger.info(f"Pinecone connection successful. Available indexes: {[idx.name for idx in indexes]}")
            
            # Test OpenAI
            logger.info("Testing OpenAI connection...")
            response = self.openai_client.embeddings.create(
                model=OPENAI_EMBEDDING_MODEL,
                input="test connection",
                dimensions=EMBEDDING_DIMENSIONS
            )
            logger.info(f"OpenAI connection successful. Embedding dimension: {len(response.data[0].embedding)}")
            
            return True
            
        except Exception as e:
            logger.error(f"Connection test failed: {str(e)}")
            return False

    def print_final_stats(self):
        """Print processing statistics"""
        print("\n" + "="*60)
        print("PARALLEL VECTOR LOADING COMPLETE")
        print("="*60)
        print(f"Total Files Found: {self.stats.total_files}")
        print(f"✓ Successfully Processed: {self.stats.processed_files}")
        print(f"✗ Failed Files: {self.stats.failed_files}")
        print(f"📄 Total Chunks Created: {self.stats.total_chunks}")
        print(f"🔗 Vectors Uploaded to Pinecone: {self.stats.uploaded_vectors}")
        print(f"⏱ Processing Time: {self.stats.processing_time:.2f} seconds")
        print(f"⚡ Parallelization: {self.max_workers} workers")
        
        if self.stats.total_chunks > 0:
            avg_chunks = self.stats.total_chunks / max(self.stats.processed_files, 1)
            print(f"📊 Average Chunks per Document: {avg_chunks:.1f}")
        
        if self.stats.processing_time > 0:
            throughput = self.stats.processed_files / self.stats.processing_time
            print(f"🚀 Throughput: {throughput:.2f} files/second")
        
        print(f"\nPinecone Index: {self.index_name}")
        print(f"Namespace: {self.namespace}")
        print("Your legal document RAG system is ready!")


def main():
    """Interactive main function"""
    try:
        # Show welcome banner
        print_banner()
        
        # Step 1: Get API keys
        openai_key, pinecone_key = get_api_keys()
        
        # Step 2: Test connections
        if not test_connections(openai_key, pinecone_key):
            print("\n❌ Connection tests failed. Please check your API keys and try again.")
            return 1
        
        # Step 3: Get processing parameters
        batch_size, chunk_size, overlap, workers = get_processing_parameters()
        
        # Step 4: Get storage parameters
        docs_path, namespace, index_name = get_storage_parameters()
        if docs_path is None:
            print("\n❌ Setup cancelled.")
            return 1
        
        # Step 5: Get embedding parameters
        model, dimensions = get_embedding_parameters()
        
        # Step 6: Show configuration summary and confirm
        config = {
            'docs_path': docs_path,
            'namespace': namespace,
            'index_name': index_name,
            'model': model,
            'dimensions': dimensions,
            'batch_size': batch_size,
            'chunk_size': chunk_size,
            'overlap': overlap,
            'workers': workers
        }
        
        if not show_configuration_summary(config):
            print("\n❌ Processing cancelled.")
            return 1
        
        # Step 7: Initialize and run the loader
        print("\n🚀 STARTING DOCUMENT PROCESSING")
        print("=" * 50)
        
        # Update global configuration
        global CHUNK_SIZE, CHUNK_OVERLAP, BATCH_SIZE, MAX_WORKERS, OPENAI_EMBEDDING_MODEL, EMBEDDING_DIMENSIONS
        CHUNK_SIZE = chunk_size
        CHUNK_OVERLAP = overlap
        BATCH_SIZE = batch_size
        MAX_WORKERS = workers
        OPENAI_EMBEDDING_MODEL = model
        EMBEDDING_DIMENSIONS = dimensions
        
        # Initialize loader
        loader = ParallelLegalDocumentVectorLoader(
            docs_directory=docs_path,
            pinecone_api_key=pinecone_key,
            openai_api_key=openai_key,
            index_name=index_name,
            namespace=namespace,
            max_workers=workers,
            embedding_model=model,
            embedding_dimensions=dimensions
        )
        
        # Process documents
        stats = loader.process_all_documents_parallel(batch_size=batch_size)
        
        # Final results
        print("\n" + "="*80)
        print("🎉 PROCESSING COMPLETE!")
        print("="*80)
        
        if stats.failed_files == 0:
            print("✅ All documents processed successfully!")
        else:
            print(f"⚠️  Processing completed with {stats.failed_files} failures")
            print("   Check the log file for details on failed documents.")
        
        print(f"\n📊 Final Statistics:")
        print(f"   📁 Total files processed: {stats.processed_files}")
        print(f"   📄 Total chunks created: {stats.total_chunks}")
        print(f"   🔗 Vectors uploaded: {stats.uploaded_vectors}")
        if stats.processing_time > 0:
            print(f"   ⏱️  Processing time: {stats.processing_time:.1f} seconds")
            print(f"   🚀 Throughput: {stats.processed_files/stats.processing_time:.2f} files/second")
        
        print(f"\n🗃️  Your documents are now available in Pinecone:")
        print(f"   Index: {index_name}")
        print(f"   Namespace: {namespace}")
        print("   Ready for semantic search and AI retrieval!")
        
        return 0
        
    except KeyboardInterrupt:
        print("\n\n⏸️  Processing interrupted by user.")
        return 1
    except Exception as e:
        print(f"\n❌ Unexpected error: {str(e)}")
        logger.error(f"Unexpected error in main: {str(e)}")
        return 1

def legacy_main():
    """Legacy command-line interface (kept for backwards compatibility)"""
    parser = argparse.ArgumentParser(description='Load legal documents into Pinecone vector store (parallel version)')
    parser.add_argument('--docs-directory', '-d',
                       default=DEFAULT_DOCS_PATH,
                       help='Directory containing legal documents to process')
    parser.add_argument('--load-documents', action='store_true',
                       help='Process and load all documents')
    parser.add_argument('--test-connection', action='store_true',
                       help='Test connections to Pinecone and OpenAI')
    parser.add_argument('--workers', type=int, default=MAX_WORKERS,
                       help='Number of parallel workers')
    parser.add_argument('--batch-size', type=int, default=BATCH_SIZE,
                       help='Batch size for vector uploads')
    parser.add_argument('--index-name', default=PINECONE_INDEX_NAME,
                       help='Pinecone index name')
    parser.add_argument('--namespace', default=PINECONE_NAMESPACE,
                       help='Pinecone namespace')
    parser.add_argument('--verbose', '-v', action='store_true',
                       help='Enable verbose logging')
    parser.add_argument('--interactive', action='store_true',
                       help='Use interactive mode (default)')
    
    args = parser.parse_args()
    
    # Use interactive mode by default, unless specific CLI args are provided
    if not any([args.load_documents, args.test_connection]) or args.interactive:
        return main()
    
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    # Get API keys from environment
    pinecone_api_key = os.getenv('PINECONE_API_KEY')
    openai_api_key = os.getenv('OPENAI_API_KEY')
    
    if not pinecone_api_key:
        print("Error: PINECONE_API_KEY environment variable not set")
        return 1
    
    if not openai_api_key:
        print("Error: OPENAI_API_KEY environment variable not set")
        return 1
    
    try:
        # Initialize loader
        loader = ParallelLegalDocumentVectorLoader(
            docs_directory=args.docs_directory,
            pinecone_api_key=pinecone_api_key,
            openai_api_key=openai_api_key,
            index_name=args.index_name,
            namespace=args.namespace,
            max_workers=args.workers
        )
        
        if args.test_connection:
            if loader.test_connections():
                print("✓ All connections successful!")
                return 0
            else:
                print("✗ Connection test failed")
                return 1
        
        elif args.load_documents:
            stats = loader.process_all_documents_parallel(batch_size=args.batch_size)
            
            if stats.failed_files == 0:
                print("\n🎉 All documents processed successfully!")
                return 0
            else:
                print(f"\n⚠ Processing completed with {stats.failed_files} failures")
                return 0
        
        else:
            print("Please specify --load-documents or --test-connection")
            print("Use --help for more options")
            return 1
    
    except KeyboardInterrupt:
        logger.info("Process interrupted by user")
        return 0
    except Exception as e:
        logger.error(f"Process failed: {str(e)}")
        return 1


if __name__ == "__main__":
    exit(main())